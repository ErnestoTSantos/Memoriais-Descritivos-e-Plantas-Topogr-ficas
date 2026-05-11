from __future__ import annotations

from datetime import datetime
import math
from pathlib import Path

import ezdxf
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from app.models.schemas import CoordinatePoint, SegmentInfo


def generate_memorial_text(
    property_name: str,
    owner_name: str,
    municipality: str,
    state: str,
    datum: str,
    coordinate_system: str,
    measurement_mode: str,
    irradiation_origin_x: float | None,
    irradiation_origin_y: float | None,
    irradiation_angle_error_seconds: float | None,
    area_m2: float,
    perimeter_m: float,
    segments: list[SegmentInfo],
) -> str:
    date_text = datetime.now().strftime("%d/%m/%Y")
    measurement_label = "Irradiacao" if measurement_mode == "irradiacao" else "Ponto a ponto"

    irradiation_origin_line = ""
    if measurement_mode == "irradiacao" and irradiation_origin_x is not None and irradiation_origin_y is not None:
        irradiation_origin_line = f"Estacao de irradiacao (X, Y): {irradiation_origin_x:.3f}, {irradiation_origin_y:.3f}\n"
    irradiation_polygon_line = ""
    if measurement_mode == "irradiacao":
        irradiation_polygon_line = "Poligonal gerada a partir de pontos irradiados.\n"
    irradiation_error_line = ""
    if measurement_mode == "irradiacao" and irradiation_angle_error_seconds is not None:
        irradiation_error_line = f"Ajuste angular aplicado (segundos): {irradiation_angle_error_seconds:.2f}\n"

    intro = (
        f"MEMORIAL DESCRITIVO DO IMOVEL {property_name.upper()}\n\n"
        f"Proprietario: {owner_name}\n"
        f"Municipio/UF: {municipality}/{state}\n"
        f"Sistema Geodesico: {datum}\n"
        f"Sistema de Coordenadas: {coordinate_system}\n"
        f"Modalidade de medicao: {measurement_label}\n"
        f"{irradiation_origin_line}"
        f"{irradiation_polygon_line}"
        f"{irradiation_error_line}"
        f"Perimetro: {perimeter_m:.2f} m\n"
        f"Area: {area_m2:.2f} m2\n\n"
        "DESCRICAO DOS LIMITES E CONFRONTACOES:\n"
    )

    if segments and segments[-1].end_vertex != segments[0].start_vertex:
        raise ValueError(
            "Memorial invalido: poligono aberto. Inclua o trecho final do ultimo vertice ao primeiro."
        )

    lines = []
    for idx, seg in enumerate(segments, start=1):
        lines.append(
            f"{idx:02d}) Do vertice {seg.start_vertex} ao vertice {seg.end_vertex}, "
            f"seguindo com rumo {seg.bearing}, distancia de {seg.distance_m:.2f} m."
        )

    footer = (
        "\n\n"
        f"Documento gerado automaticamente em {date_text}. "
        "Recomenda-se validacao tecnica por profissional habilitado conforme diretrizes do INCRA "
        "e Provimento CNJ no 65/2017."
    )

    return intro + "\n".join(lines) + footer


def _prepare_closed_points(points: list[CoordinatePoint] | None) -> list[CoordinatePoint]:
    if not points or len(points) < 3:
        return []

    closed = list(points)
    first = closed[0]
    last = closed[-1]
    if first.vertex != last.vertex or first.x != last.x or first.y != last.y:
        closed.append(first)
    return closed


def _project_points(
    points: list[CoordinatePoint],
    left: float,
    top: float,
    width: float,
    height: float,
    padding: float = 8.0,
) -> list[tuple[float, float]]:
    xs = [p.x for p in points]
    ys = [p.y for p in points]
    min_x = min(xs)
    max_x = max(xs)
    min_y = min(ys)
    max_y = max(ys)

    span_x = max_x - min_x
    span_y = max_y - min_y
    usable_w = max(width - (2 * padding), 1.0)
    usable_h = max(height - (2 * padding), 1.0)

    scale_x = usable_w / span_x if span_x > 0 else float("inf")
    scale_y = usable_h / span_y if span_y > 0 else float("inf")
    scale = min(scale_x, scale_y)
    if not math.isfinite(scale):
        scale = 1.0

    drawn_w = span_x * scale
    drawn_h = span_y * scale
    offset_x = left + padding + ((usable_w - drawn_w) / 2)
    offset_y = top + padding + ((usable_h - drawn_h) / 2)

    projected: list[tuple[float, float]] = []
    for point in points:
        x = offset_x + ((point.x - min_x) * scale if span_x > 0 else usable_w / 2)
        y_up = offset_y + ((point.y - min_y) * scale if span_y > 0 else usable_h / 2)
        y = top + height - y_up
        projected.append((x, y))
    return projected


def _draw_polygon_on_pdf(pdf: FPDF, points: list[CoordinatePoint]) -> None:
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Croqui do poligono (vertices e segmentos)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=9)
    pdf.multi_cell(
        0,
        5,
        "Representacao esquematica em coordenadas planas, sem escala cartografica de impressao.",
    )
    pdf.ln(2)

    left = 20.0
    top = 35.0
    width = 170.0
    height = 120.0
    pdf.rect(left, top, width, height)

    projected = _project_points(points, left, top, width, height)
    for index in range(len(projected) - 1):
        x1, y1 = projected[index]
        x2, y2 = projected[index + 1]
        pdf.line(x1, y1, x2, y2)

        segment_length = math.dist((points[index].x, points[index].y), (points[index + 1].x, points[index + 1].y))
        label_x = (x1 + x2) / 2
        label_y = (y1 + y2) / 2
        pdf.set_xy(label_x - 10, label_y - 2)
        pdf.set_font("Helvetica", size=7)
        pdf.cell(20, 4, f"S{index + 1}: {segment_length:.2f} m", align="C")

    pdf.set_font("Helvetica", size=8)
    for index, (x, y) in enumerate(projected[:-1]):
        pdf.ellipse(x - 1.1, y - 1.1, 2.2, 2.2)
        pdf.set_xy(x + 1.5, y - 2.2)
        pdf.cell(18, 4, points[index].vertex)

    pdf.set_xy(left, top + height + 4)
    pdf.set_font("Helvetica", size=8)
    pdf.multi_cell(0, 4, f"Total de vertices: {len(points) - 1} | Total de segmentos: {len(points) - 1}")


def _build_ascii_sketch(points: list[CoordinatePoint], width: int = 64, height: int = 20) -> str:
    if width < 8 or height < 8:
        raise ValueError("Dimensoes minimas do croqui ASCII: 8x8.")

    projected = _project_points(points, 0.0, 0.0, float(width - 1), float(height - 1), padding=1.0)
    grid_points = [(int(round(x)), int(round(y))) for x, y in projected]
    canvas = [[" " for _ in range(width)] for _ in range(height)]

    def draw_line(x0: int, y0: int, x1: int, y1: int) -> None:
        dx = abs(x1 - x0)
        dy = -abs(y1 - y0)
        sx = 1 if x0 < x1 else -1
        sy = 1 if y0 < y1 else -1
        err = dx + dy
        cx = x0
        cy = y0

        while True:
            if 0 <= cx < width and 0 <= cy < height and canvas[cy][cx] == " ":
                canvas[cy][cx] = "#"
            if cx == x1 and cy == y1:
                break
            e2 = 2 * err
            if e2 >= dy:
                err += dy
                cx += sx
            if e2 <= dx:
                err += dx
                cy += sy

    for i in range(len(grid_points) - 1):
        x0, y0 = grid_points[i]
        x1, y1 = grid_points[i + 1]
        draw_line(x0, y0, x1, y1)

    for x, y in grid_points[:-1]:
        if 0 <= x < width and 0 <= y < height:
            canvas[y][x] = "o"

    return "\n".join("".join(row).rstrip() for row in canvas)


def export_pdf(
    path: Path,
    title: str,
    memorial_text: str,
    points: list[CoordinatePoint] | None = None,
) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, title)
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, memorial_text)

    closed_points = _prepare_closed_points(points)
    if closed_points:
        _draw_polygon_on_pdf(pdf, closed_points)

    pdf.output(str(path))


def export_docx(
    path: Path,
    title: str,
    memorial_text: str,
    points: list[CoordinatePoint] | None = None,
) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in memorial_text.splitlines():
        doc.add_paragraph(line)

    closed_points = _prepare_closed_points(points)
    if closed_points:
        doc.add_page_break()
        doc.add_heading("Croqui do poligono e segmentos", level=2)
        doc.add_paragraph(
            "Representacao tecnica auxiliar em coordenadas planas (esquematica, sem escala de impressao)."
        )

        segment_table = doc.add_table(rows=1, cols=4)
        header = segment_table.rows[0].cells
        header[0].text = "Segmento"
        header[1].text = "De"
        header[2].text = "Para"
        header[3].text = "Distancia (m)"

        for idx in range(len(closed_points) - 1):
            start = closed_points[idx]
            end = closed_points[idx + 1]
            row = segment_table.add_row().cells
            row[0].text = f"S{idx + 1:02d}"
            row[1].text = start.vertex
            row[2].text = end.vertex
            row[3].text = f"{math.dist((start.x, start.y), (end.x, end.y)):.2f}"

        doc.add_paragraph()
        doc.add_paragraph("Croqui textual (ASCII):")
        sketch_paragraph = doc.add_paragraph()
        sketch_run = sketch_paragraph.add_run(_build_ascii_sketch(closed_points))
        sketch_run.font.name = "Courier New"
        sketch_run.font.size = Pt(8)

        doc.add_paragraph()
        doc.add_paragraph("Vertices:")
        for point in closed_points[:-1]:
            doc.add_paragraph(f"{point.vertex}: X={point.x:.3f} | Y={point.y:.3f}")

    doc.save(str(path))


def export_dxf(path: Path, points: list[CoordinatePoint], title: str) -> None:
    doc = ezdxf.new(setup=True)
    msp = doc.modelspace()

    poly_points = [(p.x, p.y) for p in points]
    msp.add_lwpolyline(poly_points, close=True, dxfattribs={"layer": "LIMITES"})

    for p in points:
        msp.add_circle((p.x, p.y), radius=0.4, dxfattribs={"layer": "VERTICES"})
        msp.add_text(
            p.vertex,
            dxfattribs={"height": 1.2, "layer": "ROTULOS"},
        ).set_placement((p.x + 0.6, p.y + 0.6))

    min_x = min(p.x for p in points)
    max_y = max(p.y for p in points)
    msp.add_text(
        title,
        dxfattribs={"height": 2.0, "layer": "TITULO"},
    ).set_placement((min_x, max_y + 4))

    doc.saveas(str(path))
